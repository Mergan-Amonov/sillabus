import io
from uuid import UUID
from datetime import date
from sqlalchemy.ext.asyncio import AsyncSession
from sqlalchemy import select

from app.core.exceptions import NotFoundError, ForbiddenError
from app.modules.auth.models import User, UserRole
from app.modules.syllabus.models import Syllabus, SyllabusStatus


# ─────────────────────────────────────────────────────────────────────────────
# DB helper
# ─────────────────────────────────────────────────────────────────────────────

async def get_syllabus_for_export(syllabus_id: UUID, user: User, db: AsyncSession) -> Syllabus:
    result = await db.execute(select(Syllabus).where(Syllabus.id == syllabus_id))
    syllabus = result.scalar_one_or_none()
    if not syllabus:
        raise NotFoundError("Syllabus")

    if user.role == UserRole.TEACHER and syllabus.created_by != user.id:
        raise ForbiddenError("Access denied")
    if user.role in (UserRole.REVIEWER, UserRole.UNIVERSITY_ADMIN):
        if syllabus.university_id != user.university_id:
            raise ForbiddenError("Access denied")

    if syllabus.status not in (SyllabusStatus.APPROVED, SyllabusStatus.PENDING_REVIEW):
        if user.role == UserRole.TEACHER and syllabus.created_by == user.id:
            pass  # teacher can export own draft
        else:
            raise ForbiddenError("Only approved or pending syllabuses can be exported")

    return syllabus


# ─────────────────────────────────────────────────────────────────────────────
# Shared helpers
# ─────────────────────────────────────────────────────────────────────────────

_STATUS_LABELS = {
    "approved": "TASDIQLANGAN",
    "draft": "QORALAMA",
    "pending_review": "KO'RIB CHIQILMOQDA",
    "rejected": "RAD ETILGAN",
    "archived": "ARXIVLANGAN",
}

_GRADING_LABELS = {
    "current_control": "Joriy nazorat",
    "midterm": "Oraliq imtihon",
    "final": "Yakuniy imtihon",
}

def _fmt_textbook(t: dict, idx: int) -> str:
    author = t.get("author", "")
    title = t.get("title", "")
    pub = t.get("publisher", "")
    year = t.get("year", "")
    parts = [p for p in [author, title, pub, str(year) if year else ""] if p]
    return f"[{idx}] " + ". ".join(parts) + "."

def _week_cells(w: dict) -> tuple[str, str, str, str]:
    """Return (week_no, topic, lecture, practice) handling v1 and v2 formats."""
    wn = str(w.get("week", ""))
    topic = w.get("topic", "")
    lecture = w.get("lecture_content", "") or w.get("description", "") or ""
    practice = w.get("practice_content", "") or ""
    if not practice:
        acts = w.get("activities", "")
        practice = ", ".join(acts) if isinstance(acts, list) else (acts or "")
    self_study = w.get("self_study", "") or ""
    return wn, topic, str(lecture), str(practice), str(self_study)


# ─────────────────────────────────────────────────────────────────────────────
# PDF
# ─────────────────────────────────────────────────────────────────────────────

def generate_pdf(syllabus: Syllabus) -> bytes:
    from reportlab.lib.pagesizes import A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.lib.enums import TA_CENTER, TA_LEFT
    from reportlab.lib import colors
    from reportlab.platypus import (
        SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
    )

    PAGE_W, PAGE_H = A4
    MARGIN = 2.5 * cm
    USABLE_W = PAGE_W - 2 * MARGIN

    BLUE = colors.HexColor("#2563EB")
    DARK_BLUE = colors.HexColor("#1D4ED8")
    LIGHT_GRAY = colors.HexColor("#F9FAFB")
    ROW_ALT = colors.HexColor("#F3F4F6")
    BORDER = colors.HexColor("#D1D5DB")
    TEXT_GRAY = colors.HexColor("#4B5563")

    buffer = io.BytesIO()
    styles = getSampleStyleSheet()

    # ── Styles ──────────────────────────────────────────────────────────────
    title_st = ParagraphStyle(
        "CT", fontSize=18, fontName="Helvetica-Bold",
        textColor=colors.HexColor("#111827"), alignment=TA_CENTER, spaceAfter=6,
    )
    sub_st = ParagraphStyle(
        "CS", fontSize=12, fontName="Helvetica",
        textColor=TEXT_GRAY, alignment=TA_CENTER, spaceAfter=4,
    )
    meta_line_st = ParagraphStyle(
        "ML", fontSize=10, fontName="Helvetica",
        textColor=TEXT_GRAY, alignment=TA_CENTER,
    )
    section_st = ParagraphStyle(
        "SH", fontSize=13, fontName="Helvetica-Bold",
        textColor=DARK_BLUE, spaceBefore=14, spaceAfter=4,
    )
    normal_st = ParagraphStyle(
        "CN", fontSize=10, fontName="Helvetica", leading=14, spaceAfter=4,
    )
    bullet_st = ParagraphStyle(
        "CB", fontSize=10, fontName="Helvetica", leading=14,
        leftIndent=14, spaceAfter=3,
    )
    subhead_st = ParagraphStyle(
        "CSH", fontSize=10, fontName="Helvetica-Bold", spaceAfter=4,
    )

    # ── Footer ───────────────────────────────────────────────────────────────
    today = date.today().strftime("%d.%m.%Y")
    footer_text = f"Silabuys | {syllabus.course_code} | {today}"

    def _footer(canvas, doc):
        canvas.saveState()
        canvas.setFont("Helvetica", 8)
        canvas.setFillColor(colors.HexColor("#9CA3AF"))
        canvas.drawString(MARGIN, MARGIN * 0.6, footer_text)
        canvas.drawRightString(PAGE_W - MARGIN, MARGIN * 0.6, f"Bet {doc.page}")
        canvas.restoreState()

    # ── Helpers ──────────────────────────────────────────────────────────────
    def hr():
        return HRFlowable(width="100%", thickness=0.5, color=BORDER, spaceAfter=4)

    def section(title: str):
        story.append(Spacer(1, 0.5 * cm))
        story.append(Paragraph(title, section_st))
        story.append(hr())

    def overview_table(rows: list[list[str]]) -> Table:
        t = Table(rows, colWidths=[5 * cm, USABLE_W - 5 * cm])
        t.setStyle(TableStyle([
            ("FONTNAME", (0, 0), (0, -1), "Helvetica-Bold"),
            ("FONTNAME", (1, 0), (1, -1), "Helvetica"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
            ("ROWBACKGROUNDS", (0, 0), (-1, -1), [LIGHT_GRAY, colors.white]),
            ("GRID", (0, 0), (-1, -1), 0.3, BORDER),
        ]))
        return t

    story: list = []

    # ── 1. Sarlavha sahifasi ─────────────────────────────────────────────────
    logo_table = Table([["SILABUYS"]], colWidths=[USABLE_W])
    logo_table.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), BLUE),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 14),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 12),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 12),
    ]))
    story.append(logo_table)
    story.append(Spacer(1, 0.8 * cm))
    story.append(Paragraph(syllabus.title, title_st))

    dept_parts = [p for p in [syllabus.department, syllabus.faculty] if p]
    if dept_parts:
        story.append(Paragraph(" · ".join(dept_parts), sub_st))

    story.append(Spacer(1, 0.4 * cm))

    # Status badge
    status_val = syllabus.status.value if hasattr(syllabus.status, "value") else str(syllabus.status)
    status_color_map = {
        "approved": "#16A34A", "draft": "#6B7280",
        "pending_review": "#D97706", "rejected": "#DC2626", "archived": "#4B5563",
    }
    badge_color = colors.HexColor(status_color_map.get(status_val, "#6B7280"))
    badge_label = _STATUS_LABELS.get(status_val, status_val.upper())
    badge = Table([[badge_label]], colWidths=[3.5 * cm])
    badge.setStyle(TableStyle([
        ("BACKGROUND", (0, 0), (-1, -1), badge_color),
        ("TEXTCOLOR", (0, 0), (-1, -1), colors.white),
        ("FONTNAME", (0, 0), (-1, -1), "Helvetica-Bold"),
        ("FONTSIZE", (0, 0), (-1, -1), 9),
        ("ALIGN", (0, 0), (-1, -1), "CENTER"),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]))
    story.append(Table([[badge]], colWidths=[USABLE_W],
                       style=TableStyle([("ALIGN", (0, 0), (-1, -1), "CENTER")])))

    story.append(Spacer(1, 0.5 * cm))
    mini_parts = [f"Kurs kodi: {syllabus.course_code}", f"Kredit soat: {syllabus.credit_hours}"]
    if syllabus.semester:
        mini_parts.append(f"Semestr: {syllabus.semester}")
    if syllabus.academic_year:
        mini_parts.append(f"O'quv yili: {syllabus.academic_year}")
    story.append(Paragraph("  |  ".join(mini_parts), meta_line_st))
    story.append(Spacer(1, 0.3 * cm))
    story.append(Paragraph("O'qituvchi: ____________________________",
                           ParagraphStyle("T", fontSize=10, fontName="Helvetica",
                                          textColor=TEXT_GRAY, alignment=TA_CENTER)))

    # ── 2. Kurs haqida ───────────────────────────────────────────────────────
    section("KURS HAQIDA")
    rows: list[list[str]] = [
        ["Kurs kodi", syllabus.course_code],
        ["Kredit soat", str(syllabus.credit_hours)],
    ]
    if syllabus.semester:
        rows.append(["Semestr", str(syllabus.semester)])
    if syllabus.academic_year:
        rows.append(["O'quv yili", syllabus.academic_year])
    if syllabus.department:
        rows.append(["Kafedra", syllabus.department])
    if syllabus.faculty:
        rows.append(["Fakultet", syllabus.faculty])
    if syllabus.language:
        rows.append(["O'qitish tili", syllabus.language])
    if syllabus.prerequisites:
        rows.append(["Prerekvizitlar", syllabus.prerequisites])
    story.append(overview_table(rows))

    # ── 3. Soatlar taqsimoti ─────────────────────────────────────────────────
    lh = syllabus.lecture_hours
    ph = syllabus.practice_hours
    lbh = syllabus.lab_hours
    ssh = syllabus.self_study_hours
    if any(x is not None for x in [lh, ph, lbh, ssh]):
        section("SOATLAR TAQSIMOTI")
        total = (lh or 0) + (ph or 0) + (lbh or 0) + (ssh or 0)
        h_table = Table(
            [["Ma'ruza", "Amaliy", "Laboratoriya", "Mustaqil ish", "Jami"],
             [str(lh or 0), str(ph or 0), str(lbh or 0), str(ssh or 0), str(total)]],
            colWidths=[USABLE_W / 5] * 5,
        )
        h_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ALIGN", (0, 0), (-1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
            ("BACKGROUND", (4, 1), (4, 1), colors.HexColor("#DBEAFE")),
            ("FONTNAME", (4, 1), (4, 1), "Helvetica-Bold"),
            ("TOPPADDING", (0, 0), (-1, -1), 6),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 6),
        ]))
        story.append(h_table)

    # ── 4. Kurs tavsifi ──────────────────────────────────────────────────────
    if syllabus.description:
        section("KURS TAVSIFI")
        story.append(Paragraph(syllabus.description, normal_st))

    # ── 5. O'quv natijalari ──────────────────────────────────────────────────
    if syllabus.learning_outcomes:
        section("O'QUV NATIJALARI")
        for item in syllabus.learning_outcomes:
            story.append(Paragraph(f"• {item}", bullet_st))

    # ── 6. Kompetensiyalar ───────────────────────────────────────────────────
    if syllabus.competencies:
        section("KOMPETENSIYALAR")
        for item in syllabus.competencies:
            story.append(Paragraph(f"• {item}", bullet_st))

    # ── 7. Haftalik jadval ───────────────────────────────────────────────────
    content = syllabus.content or {}
    weeks = content.get("weeks", [])
    if weeks:
        section("HAFTALIK JADVAL")
        col_w = USABLE_W / 5
        w_data = [["#", "Mavzu", "Ma'ruza", "Amaliy", "Mustaqil ish"]]
        for w in weeks:
            wn, topic, lecture, practice, self_study = _week_cells(w)
            w_data.append([wn, topic[:80], lecture[:80], practice[:80], self_study[:80]])

        w_styles = [
            ("BACKGROUND", (0, 0), (-1, 0), BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 8),
            ("GRID", (0, 0), (-1, -1), 0.3, BORDER),
            ("ALIGN", (0, 0), (0, -1), "CENTER"),
            ("VALIGN", (0, 0), (-1, -1), "TOP"),
            ("TOPPADDING", (0, 0), (-1, -1), 4),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
        ]
        for i in range(1, len(w_data)):
            bg = LIGHT_GRAY if i % 2 == 0 else colors.white
            w_styles.append(("BACKGROUND", (0, i), (-1, i), bg))

        w_table = Table(w_data, colWidths=[1.0 * cm, 3.5 * cm, col_w, col_w, col_w])
        w_table.setStyle(TableStyle(w_styles))
        story.append(w_table)

    # ── 8. Baholash tizimi ───────────────────────────────────────────────────
    grading = syllabus.grading_policy
    if grading:
        section("BAHOLASH TIZIMI")
        g_data = [["Baholash turi", "Ulush"]]
        for key, val in grading.items():
            label = _GRADING_LABELS.get(key, key.replace("_", " ").title())
            g_data.append([label, f"{val}%"])

        g_table = Table(g_data, colWidths=[USABLE_W * 0.72, USABLE_W * 0.28])
        g_table.setStyle(TableStyle([
            ("BACKGROUND", (0, 0), (-1, 0), BLUE),
            ("TEXTCOLOR", (0, 0), (-1, 0), colors.white),
            ("FONTNAME", (0, 0), (-1, 0), "Helvetica-Bold"),
            ("FONTNAME", (1, 1), (1, -1), "Helvetica-Bold"),
            ("FONTSIZE", (0, 0), (-1, -1), 10),
            ("ALIGN", (1, 0), (1, -1), "CENTER"),
            ("GRID", (0, 0), (-1, -1), 0.5, BORDER),
            ("ROWBACKGROUNDS", (0, 1), (-1, -1), [colors.white, LIGHT_GRAY]),
            ("TOPPADDING", (0, 0), (-1, -1), 5),
            ("BOTTOMPADDING", (0, 0), (-1, -1), 5),
        ]))
        story.append(g_table)

        if syllabus.passing_grade is not None:
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph(
                f"O'tish bali: <b>{syllabus.passing_grade}</b> ball", normal_st
            ))

    # ── 9. Davomat talablari ─────────────────────────────────────────────────
    if syllabus.attendance_policy:
        section("DAVOMATGA TALABLAR")
        story.append(Paragraph(syllabus.attendance_policy, normal_st))

    # ── 10. Adabiyotlar ──────────────────────────────────────────────────────
    textbooks = syllabus.textbooks or []
    if textbooks:
        section("ADABIYOTLAR")
        required = [t for t in textbooks if t.get("required")]
        optional = [t for t in textbooks if not t.get("required")]

        if required:
            story.append(Paragraph("Asosiy adabiyotlar:", subhead_st))
            for i, t in enumerate(required, 1):
                story.append(Paragraph(_fmt_textbook(t, i), bullet_st))

        if optional:
            story.append(Spacer(1, 0.2 * cm))
            story.append(Paragraph("Qo'shimcha adabiyotlar:", subhead_st))
            for i, t in enumerate(optional, 1):
                story.append(Paragraph(_fmt_textbook(t, i), bullet_st))

    # ── 11. Onlayn resurslar ─────────────────────────────────────────────────
    online_resources = syllabus.online_resources or []
    if online_resources:
        section("ONLAYN RESURSLAR")
        for r in online_resources:
            name = r.get("name", "")
            url = r.get("url", "")
            desc = r.get("description", "")
            line = f"• {name}"
            if url:
                line += f" — {url}"
            if desc:
                line += f" ({desc})"
            story.append(Paragraph(line, bullet_st))

    doc = SimpleDocTemplate(
        buffer,
        pagesize=A4,
        leftMargin=MARGIN,
        rightMargin=MARGIN,
        topMargin=MARGIN,
        bottomMargin=MARGIN,
    )
    doc.build(story, onFirstPage=_footer, onLaterPages=_footer)
    return buffer.getvalue()


# ─────────────────────────────────────────────────────────────────────────────
# DOCX
# ─────────────────────────────────────────────────────────────────────────────

def _docx_set_cell_bg(cell, hex_color: str) -> None:
    """Set table cell background color using low-level XML."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color.lstrip("#"))
    tcPr.append(shd)


def _docx_header_row(row, texts: list[str], bg: str = "2563EB") -> None:
    """Style a table header row: blue background, white bold text."""
    from docx.shared import RGBColor, Pt

    for i, cell in enumerate(row.cells):
        cell.text = texts[i] if i < len(texts) else ""
        _docx_set_cell_bg(cell, bg)
        run = cell.paragraphs[0].runs
        if run:
            run[0].bold = True
            run[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run[0].font.size = Pt(10)


def _docx_alt_row(row, texts: list[str], even: bool) -> None:
    """Fill a data row with optional alternating background."""
    from docx.shared import Pt

    bg = "F3F4F6" if even else "FFFFFF"
    for i, cell in enumerate(row.cells):
        cell.text = texts[i] if i < len(texts) else ""
        _docx_set_cell_bg(cell, bg)
        if cell.paragraphs[0].runs:
            cell.paragraphs[0].runs[0].font.size = Pt(9)


def generate_docx(syllabus: Syllabus) -> bytes:
    from docx import Document
    from docx.shared import Pt, RGBColor, Cm
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    doc = Document()

    # ── Page margins ─────────────────────────────────────────────────────────
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # ── Helper: section heading ───────────────────────────────────────────────
    def add_section(title: str) -> None:
        p = doc.add_heading(title, level=1)
        p.runs[0].font.color.rgb = RGBColor(0x1D, 0x4E, 0xD8)
        p.runs[0].font.size = Pt(13)

    def add_bullet(text: str) -> None:
        doc.add_paragraph(text, style="List Bullet")

    def add_numbered(text: str) -> None:
        doc.add_paragraph(text, style="List Number")

    # ── 1. Sarlavha sahifasi ─────────────────────────────────────────────────
    title_p = doc.add_heading(syllabus.title, level=0)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if title_p.runs:
        title_p.runs[0].font.size = Pt(18)

    dept_parts = [p for p in [syllabus.department, syllabus.faculty] if p]
    if dept_parts:
        sub = doc.add_paragraph(" · ".join(dept_parts))
        sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
        if sub.runs:
            sub.runs[0].font.size = Pt(12)
            sub.runs[0].font.color.rgb = RGBColor(0x4B, 0x55, 0x63)

    status_val = syllabus.status.value if hasattr(syllabus.status, "value") else str(syllabus.status)
    status_p = doc.add_paragraph(_STATUS_LABELS.get(status_val, status_val.upper()))
    status_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if status_p.runs:
        status_p.runs[0].bold = True

    mini_parts = [f"Kurs kodi: {syllabus.course_code}", f"Kredit soat: {syllabus.credit_hours}"]
    if syllabus.semester:
        mini_parts.append(f"Semestr: {syllabus.semester}")
    if syllabus.academic_year:
        mini_parts.append(f"O'quv yili: {syllabus.academic_year}")
    meta_p = doc.add_paragraph("  |  ".join(mini_parts))
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    teacher_p = doc.add_paragraph("O'qituvchi: ____________________________")
    teacher_p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # ── 2. Kurs haqida ───────────────────────────────────────────────────────
    add_section("KURS HAQIDA")
    meta_rows: list[tuple[str, str]] = [
        ("Kurs kodi", syllabus.course_code),
        ("Kredit soat", str(syllabus.credit_hours)),
    ]
    if syllabus.semester:
        meta_rows.append(("Semestr", str(syllabus.semester)))
    if syllabus.academic_year:
        meta_rows.append(("O'quv yili", syllabus.academic_year))
    if syllabus.department:
        meta_rows.append(("Kafedra", syllabus.department))
    if syllabus.faculty:
        meta_rows.append(("Fakultet", syllabus.faculty))
    if syllabus.language:
        meta_rows.append(("O'qitish tili", syllabus.language))
    if syllabus.prerequisites:
        meta_rows.append(("Prerekvizitlar", syllabus.prerequisites))

    meta_tbl = doc.add_table(rows=len(meta_rows), cols=2)
    meta_tbl.style = "Table Grid"
    for i, (label, value) in enumerate(meta_rows):
        cells = meta_tbl.rows[i].cells
        cells[0].text = label
        cells[1].text = value
        if cells[0].paragraphs[0].runs:
            cells[0].paragraphs[0].runs[0].bold = True
        bg = "F9FAFB" if i % 2 == 0 else "FFFFFF"
        _docx_set_cell_bg(cells[0], bg)
        _docx_set_cell_bg(cells[1], bg)
    doc.add_paragraph()

    # ── 3. Soatlar taqsimoti ─────────────────────────────────────────────────
    lh = syllabus.lecture_hours
    ph = syllabus.practice_hours
    lbh = syllabus.lab_hours
    ssh = syllabus.self_study_hours
    if any(x is not None for x in [lh, ph, lbh, ssh]):
        add_section("SOATLAR TAQSIMOTI")
        total = (lh or 0) + (ph or 0) + (lbh or 0) + (ssh or 0)
        h_tbl = doc.add_table(rows=2, cols=5)
        h_tbl.style = "Table Grid"
        _docx_header_row(h_tbl.rows[0],
                         ["Ma'ruza", "Amaliy", "Laboratoriya", "Mustaqil ish", "Jami"])
        _docx_alt_row(h_tbl.rows[1],
                      [str(lh or 0), str(ph or 0), str(lbh or 0), str(ssh or 0), str(total)],
                      even=False)
        doc.add_paragraph()

    # ── 4. Kurs tavsifi ──────────────────────────────────────────────────────
    if syllabus.description:
        add_section("KURS TAVSIFI")
        doc.add_paragraph(syllabus.description)

    # ── 5. O'quv natijalari ──────────────────────────────────────────────────
    if syllabus.learning_outcomes:
        add_section("O'QUV NATIJALARI")
        for item in syllabus.learning_outcomes:
            add_bullet(str(item))

    # ── 6. Kompetensiyalar ───────────────────────────────────────────────────
    if syllabus.competencies:
        add_section("KOMPETENSIYALAR")
        for item in syllabus.competencies:
            add_bullet(str(item))

    # ── 7. Haftalik jadval ───────────────────────────────────────────────────
    content = syllabus.content or {}
    weeks = content.get("weeks", [])
    if weeks:
        add_section("HAFTALIK JADVAL")
        sched_tbl = doc.add_table(rows=1 + len(weeks), cols=5)
        sched_tbl.style = "Table Grid"
        _docx_header_row(sched_tbl.rows[0],
                         ["#", "Mavzu", "Ma'ruza", "Amaliy", "Mustaqil ish"])
        for i, w in enumerate(weeks):
            wn, topic, lecture, practice, self_study = _week_cells(w)
            _docx_alt_row(sched_tbl.rows[i + 1],
                          [wn, topic, lecture, practice, self_study],
                          even=(i % 2 == 1))
        doc.add_paragraph()

    # ── 8. Baholash tizimi ───────────────────────────────────────────────────
    grading = syllabus.grading_policy
    if grading:
        add_section("BAHOLASH TIZIMI")
        g_rows = [(
            _GRADING_LABELS.get(k, k.replace("_", " ").title()),
            f"{v}%"
        ) for k, v in grading.items()]
        g_tbl = doc.add_table(rows=1 + len(g_rows), cols=2)
        g_tbl.style = "Table Grid"
        _docx_header_row(g_tbl.rows[0], ["Baholash turi", "Ulush"])
        for i, (label, val) in enumerate(g_rows):
            _docx_alt_row(g_tbl.rows[i + 1], [label, val], even=(i % 2 == 1))

        if syllabus.passing_grade is not None:
            p = doc.add_paragraph()
            run = p.add_run(f"O'tish bali: {syllabus.passing_grade} ball")
            run.bold = True
        doc.add_paragraph()

    # ── 9. Davomat talablari ─────────────────────────────────────────────────
    if syllabus.attendance_policy:
        add_section("DAVOMATGA TALABLAR")
        doc.add_paragraph(syllabus.attendance_policy)

    # ── 10. Adabiyotlar ──────────────────────────────────────────────────────
    textbooks = syllabus.textbooks or []
    if textbooks:
        add_section("ADABIYOTLAR")
        required = [t for t in textbooks if t.get("required")]
        optional = [t for t in textbooks if not t.get("required")]

        if required:
            p = doc.add_paragraph("Asosiy adabiyotlar:")
            p.runs[0].bold = True
            for i, t in enumerate(required, 1):
                add_numbered(_fmt_textbook(t, i))

        if optional:
            p = doc.add_paragraph("Qo'shimcha adabiyotlar:")
            p.runs[0].bold = True
            for i, t in enumerate(optional, 1):
                add_bullet(_fmt_textbook(t, i))

    # ── 11. Onlayn resurslar ─────────────────────────────────────────────────
    online_resources = syllabus.online_resources or []
    if online_resources:
        add_section("ONLAYN RESURSLAR")
        for r in online_resources:
            name = r.get("name", "")
            url = r.get("url", "")
            desc = r.get("description", "")
            parts = [name]
            if url:
                parts.append(url)
            if desc:
                parts.append(f"({desc})")
            add_bullet(" — ".join(parts))

    buffer = io.BytesIO()
    doc.save(buffer)
    return buffer.getvalue()
